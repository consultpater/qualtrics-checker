"""Parse a questionnaire spec doc (PDF or DOCX) into a list of expected questions."""
from __future__ import annotations

import io
import re
from typing import List

from pypdf import PdfReader
from docx import Document

from .models import SpecQuestion


# A "question" line typically starts with a number/label like: "1.", "1)", "Q1.", "Q1)", "1a.", etc.
# and ends with "?" or a period. We're lenient because spec docs vary.
QUESTION_PREFIX = re.compile(r"^\s*(Q?\d+[a-zA-Z]?)[\.\)\:]\s*(.+)$")
# Fallback: a line that ends in "?" is likely a question.
QUESTION_MARK = re.compile(r"\?\s*$")


def _text_from_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n".join(parts)


def _text_from_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    lines: List[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            lines.append(p.text)
    # Include tables too — common in questionnaire specs.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        lines.append(p.text)
    return "\n".join(lines)


def extract_spec_text(filename: str, data: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return _text_from_pdf(data)
    if name.endswith(".docx"):
        return _text_from_docx(data)
    # Best-effort plain text fallback.
    try:
        return data.decode("utf-8", errors="ignore")
    except Exception:
        return ""


def parse_spec(filename: str, data: bytes) -> List[SpecQuestion]:
    raw = extract_spec_text(filename, data)
    questions: List[SpecQuestion] = []

    # Normalize whitespace but preserve line breaks.
    lines = [ln.strip() for ln in raw.splitlines()]
    buffer: List[str] = []
    current: SpecQuestion | None = None

    def flush():
        nonlocal current
        if current is not None:
            # Join any continuation lines into the text.
            if buffer:
                current.text = (current.text + " " + " ".join(buffer)).strip()
            # Trim noise.
            current.text = re.sub(r"\s+", " ", current.text).strip()
            if current.text:
                questions.append(current)
        current = None
        buffer.clear()

    for ln in lines:
        if not ln:
            flush()
            continue
        m = QUESTION_PREFIX.match(ln)
        if m:
            flush()
            current = SpecQuestion(number=m.group(1), text=m.group(2).strip(), raw_line=ln)
        elif QUESTION_MARK.search(ln) and current is None:
            # Bare question line with no number prefix.
            flush()
            current = SpecQuestion(number=None, text=ln, raw_line=ln)
        else:
            if current is not None:
                # Continuation of a question's text (wrapped line).
                # Stop appending if we hit answer-option-looking lines.
                if re.match(r"^\s*[-*•▢☐○]|^\s*\(?\s*[a-e]\s*[\)\.]\s+|^\s*\d+\s*[\)\.]\s+", ln):
                    flush()
                else:
                    buffer.append(ln)
    flush()

    # De-dupe exact text duplicates while preserving order.
    seen = set()
    unique: List[SpecQuestion] = []
    for q in questions:
        key = q.text.lower()
        if key in seen:
            continue
        seen.add(key)
        unique.append(q)
    return unique
